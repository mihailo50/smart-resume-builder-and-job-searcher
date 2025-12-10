"""
Premium Resume Exporter - Generates stunning PDF and DOCX files.

Uses the new PremiumResumePDFGenerator for beautiful template-based exports.
"""
import logging
import io
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Try to import the premium PDF generator
try:
    from config.services.resume_pdf_generator import PremiumResumePDFGenerator
    PREMIUM_GENERATOR_AVAILABLE = True
except ImportError:
    PREMIUM_GENERATOR_AVAILABLE = False
    logger.warning("Premium PDF generator not available. Using fallback.")

# Try to import reportlab for fallback
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class ResumeExporter:
    """
    Premium resume exporter with template support.
    """
    
    def __init__(self):
        """Initialize the resume exporter."""
        self.premium_generator = None
        
        if PREMIUM_GENERATOR_AVAILABLE:
            try:
                self.premium_generator = PremiumResumePDFGenerator()
                logger.info("Premium PDF generator initialized")
            except Exception as e:
                logger.warning(f"Could not initialize premium PDF generator: {e}")
                self.premium_generator = None
    
    def export_to_pdf(
        self,
        resume_data: Dict[str, Any],
        template_name: str = 'modern-indigo',
        font_combination: str = 'modern',
        ats_mode: bool = False,
        photo_data: Optional[bytes] = None,
        photo_url: Optional[str] = None
    ) -> bytes:
        """
        Export resume data to PDF format using premium templates.
        
        Args:
            resume_data: Dictionary containing resume data
            template_name: Name of the template to use
            font_combination: Font combination (modern, classic, creative)
            ats_mode: If True, generate ATS-friendly version
            photo_data: Optional photo image bytes
            
        Returns:
            PDF file content as bytes
        """
        # Try premium generator first
        if self.premium_generator:
            try:
                pdf_bytes, _ = self.premium_generator.generate_pdf(
                    resume_data=resume_data,
                    template_name=template_name,
                    font_combination=font_combination,
                    ats_mode=ats_mode,
                    photo_data=photo_data,
                    photo_url=photo_url
                )
                return pdf_bytes
            except Exception as e:
                logger.error(f"Error with premium PDF generator: {e}")
                # Fall through to fallback
        
        # Fallback to basic reportlab if available
        if REPORTLAB_AVAILABLE:
            logger.warning("Using fallback reportlab PDF generation")
            return self._export_to_pdf_fallback(resume_data)
        
        raise ImportError(
            "PDF export requires either WeasyPrint (for premium templates) or reportlab. "
            "Install with: pip install weasyprint OR pip install reportlab"
        )
    
    def _export_to_pdf_fallback(self, resume_data: Dict[str, Any]) -> bytes:
        """Fallback PDF generation using reportlab (basic)."""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=0.75*inch, leftMargin=0.75*inch)
            story = []
            styles = getSampleStyleSheet()
            
            # Name
            if resume_data.get('full_name'):
                story.append(Paragraph(resume_data['full_name'], styles['Title']))
                story.append(Spacer(1, 0.2*inch))
            
            summary_text = resume_data.get('optimized_summary') or resume_data.get('summary')
            if summary_text:
                story.append(Paragraph('<b>Professional Summary</b>', styles['Heading2']))
                story.append(Paragraph(summary_text, styles['Normal']))
                story.append(Spacer(1, 0.3*inch))
            
            # Experiences
            if resume_data.get('experiences'):
                story.append(Paragraph('<b>Work Experience</b>', styles['Heading2']))
                for exp in resume_data['experiences']:
                    exp_text = f"<b>{exp.get('position', '')}</b> - {exp.get('company', '')}"
                    story.append(Paragraph(exp_text, styles['Normal']))
                    if exp.get('description'):
                        story.append(Paragraph(exp['description'], styles['Normal']))
                    story.append(Spacer(1, 0.2*inch))
                story.append(Spacer(1, 0.3*inch))
            
            doc.build(story)
            return buffer.getvalue()
        except Exception as e:
            logger.error(f"Error with fallback PDF generation: {e}")
            raise
    
    def export_to_docx(
        self,
        resume_data: Dict[str, Any],
        template_name: Optional[str] = None
    ) -> bytes:
        """
        Export resume data to DOCX format.
        
        Args:
            resume_data: Dictionary containing resume data
            template_name: Optional template name (for future template support in DOCX)
            
        Returns:
            DOCX file content as bytes
        """
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Name
            if resume_data.get('full_name'):
                name_para = doc.add_paragraph(resume_data['full_name'])
                name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                name_run = name_para.runs[0]
                name_run.font.size = Pt(20)
                name_run.font.bold = True
                doc.add_paragraph()
            
            summary_text = resume_data.get('optimized_summary') or resume_data.get('summary')
            if summary_text:
                doc.add_heading('Professional Summary', level=2)
                doc.add_paragraph(summary_text)
                doc.add_paragraph()
            
            # Experiences
            if resume_data.get('experiences'):
                doc.add_heading('Work Experience', level=2)
                for exp in resume_data['experiences']:
                    exp_para = doc.add_paragraph()
                    exp_run = exp_para.add_run(f"{exp.get('position', '')}")
                    exp_run.bold = True
                    exp_run.font.size = Pt(12)
                    if exp.get('company'):
                        exp_para.add_run(f" - {exp['company']}")
                    
                    if exp.get('description'):
                        doc.add_paragraph(exp['description'], style='List Bullet')
                    doc.add_paragraph()
            
            # Education
            if resume_data.get('educations'):
                doc.add_heading('Education', level=2)
                for edu in resume_data['educations']:
                    edu_para = doc.add_paragraph()
                    edu_text = edu.get('degree', '')
                    if edu.get('field_of_study'):
                        edu_text += f" in {edu['field_of_study']}"
                    if edu.get('institution'):
                        edu_text += f" - {edu['institution']}"
                    edu_run = edu_para.add_run(edu_text)
                    edu_run.bold = True
                    doc.add_paragraph()
            
            # Skills
            if resume_data.get('skills'):
                doc.add_heading('Skills', level=2)
                skills_by_category = {}
                for skill in resume_data['skills']:
                    category = skill.get('category', 'General')
                    if category not in skills_by_category:
                        skills_by_category[category] = []
                    skills_by_category[category].append(skill.get('name', ''))
                
                for category, skill_names in skills_by_category.items():
                    para = doc.add_paragraph()
                    para.add_run(f"{category}: ").bold = True
                    para.add_run(', '.join(skill_names))

                doc.add_paragraph()

            # Languages
            if resume_data.get('languages'):
                doc.add_heading('Languages', level=2)
                for lang in resume_data['languages']:
                    para = doc.add_paragraph()
                    para.add_run(lang.get('name', ''))
                    if lang.get('proficiency'):
                        para.add_run(f" ({lang['proficiency']})")
                doc.add_paragraph()

            # Interests
            if resume_data.get('interests'):
                doc.add_heading('Interests', level=2)
                interests_list = [interest.get('name', '') for interest in resume_data['interests']]
                if interests_list:
                    para = doc.add_paragraph()
                    para.add_run(', '.join(interests_list))
                doc.add_paragraph()

            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()
        
        except Exception as e:
            logger.error(f"Error exporting resume to DOCX: {e}")
            raise
    
    def get_html_preview(
        self,
        resume_data: Dict[str, Any],
        template_name: str = 'modern-indigo',
        font_combination: str = 'modern'
    ) -> str:
        """
        Get HTML preview of the resume (for frontend preview).
        
        Args:
            resume_data: Dictionary containing resume data
            template_name: Name of the template
            font_combination: Font combination
            
        Returns:
            HTML string
        """
        if self.premium_generator:
            try:
                return self.premium_generator.get_template_preview_html(
                    resume_data=resume_data,
                    template_name=template_name,
                    font_combination=font_combination
                )
            except Exception as e:
                logger.error(f"Error getting HTML preview: {e}")
        
        # Fallback
        return f"<html><body><h1>{resume_data.get('full_name', 'Resume')}</h1></body></html>"
