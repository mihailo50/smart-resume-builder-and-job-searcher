"""
Resume file parser for extracting text and data from PDF and DOCX files.
"""
import logging
import io
from typing import Dict, Any, Optional, List
from pypdf import PdfReader
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


class ResumeParser:
    """
    Service for parsing resume files (PDF and DOCX) and extracting text/data.
    """
    
    def __init__(self):
        """Initialize the resume parser."""
        pass
    
    def parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """
        Parse a PDF file and extract text.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Dict with extracted text and metadata
            
        Raises:
            Exception: If parsing fails
        """
        try:
            # Create PDF reader from bytes
            pdf_io = io.BytesIO(file_content)
            reader = PdfReader(pdf_io)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, start=1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from PDF page {page_num}: {e}")
                    continue
            
            full_text = '\n\n'.join(text_parts)
            
            # Extract metadata
            metadata = {}
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'producer': reader.metadata.get('/Producer', ''),
                }
            
            return {
                'text': full_text.strip(),
                'metadata': metadata,
                'page_count': len(reader.pages),
                'format': 'pdf',
                'success': True
            }
        
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            return {
                'text': '',
                'metadata': {},
                'page_count': 0,
                'format': 'pdf',
                'success': False,
                'error': str(e)
            }
    
    def parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """
        Parse a DOCX file and extract text.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Dict with extracted text and metadata
            
        Raises:
            Exception: If parsing fails
        """
        try:
            # Create DOCX document from bytes
            docx_io = io.BytesIO(file_content)
            doc = Document(docx_io)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        table_texts.append(' | '.join(row_texts))
            
            # Combine all text
            all_text = '\n\n'.join(paragraphs)
            if table_texts:
                all_text += '\n\n--- Tables ---\n\n' + '\n'.join(table_texts)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
            
            return {
                'text': all_text.strip(),
                'metadata': metadata,
                'paragraph_count': len(paragraphs),
                'table_count': len(doc.tables),
                'format': 'docx',
                'success': True
            }
        
        except PackageNotFoundError:
            logger.error("Invalid DOCX file: package not found")
            return {
                'text': '',
                'metadata': {},
                'paragraph_count': 0,
                'table_count': 0,
                'format': 'docx',
                'success': False,
                'error': 'Invalid DOCX file format'
            }
        
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            return {
                'text': '',
                'metadata': {},
                'paragraph_count': 0,
                'table_count': 0,
                'format': 'docx',
                'success': False,
                'error': str(e)
            }
    
    def parse_file(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse a resume file (PDF or DOCX) based on file extension or content type.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            content_type: Optional MIME type
            
        Returns:
            Dict with extracted text and metadata
        """
        # Determine file type
        file_ext = filename.split('.')[-1].lower() if '.' in filename else ''
        
        # Check content type if provided
        if content_type:
            if 'pdf' in content_type.lower():
                file_ext = 'pdf'
            elif 'wordprocessingml' in content_type.lower() or 'msword' in content_type.lower():
                file_ext = 'docx'
        
        # Parse based on file extension
        if file_ext == 'pdf':
            result = self.parse_pdf(file_content)
        elif file_ext in ['docx', 'doc']:
            result = self.parse_docx(file_content)
        else:
            # Try to detect by content
            if file_content.startswith(b'%PDF'):
                result = self.parse_pdf(file_content)
            elif file_content.startswith(b'PK'):  # DOCX is a ZIP file
                result = self.parse_docx(file_content)
            else:
                return {
                    'text': '',
                    'metadata': {},
                    'format': 'unknown',
                    'success': False,
                    'error': f'Unsupported file format: {file_ext or "unknown"}'
                }
        
        # Add filename to result
        result['filename'] = filename
        
        return result
    
    def extract_structured_data(self, parsed_text: str) -> Dict[str, Any]:
        """
        Extract structured data from parsed resume text.
        This is a basic implementation - can be enhanced with NLP/AI.
        
        Args:
            parsed_text: Extracted text from resume
            
        Returns:
            Dict with structured resume data
        """
        lines = [line.strip() for line in parsed_text.split('\n') if line.strip()]
        
        # Basic extraction (can be enhanced)
        structured_data = {
            'full_name': '',
            'email': '',
            'phone': '',
            'summary': '',
            'experiences': [],
            'educations': [],
            'skills': []
        }
        
        # Extract email
        import re
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, parsed_text)
        if emails:
            structured_data['email'] = emails[0]
        
        # Extract phone (basic pattern)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, parsed_text)
        if phones:
            structured_data['phone'] = ''.join(phones[0]) if phones[0] else ''
        
        # Extract first line as potential name
        if lines:
            structured_data['full_name'] = lines[0][:100]  # Limit length
        
        # Extract summary (first few paragraphs)
        if len(lines) > 2:
            summary_lines = [line for line in lines[1:5] if len(line) > 20]
            structured_data['summary'] = ' '.join(summary_lines[:3])[:500]
        
        return structured_data







